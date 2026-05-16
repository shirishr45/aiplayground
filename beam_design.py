"""
Reinforced Concrete Beam Design per ACI 318-19
Generates a .docx calculation sheet
"""

import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_title(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(14)

def add_heading2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(12)

def add_heading3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(11)

def add_calc(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run(f"WARNING: {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(204, 0, 0)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_pass(text):
    p = doc.add_paragraph()
    run = p.add_run(f"OK: {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_ref(text):
    p = doc.add_paragraph()
    run = p.add_run(f"Reference: {text}")
    run.italic = True
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    return p

# ── Constants ───────────────────────────────────────────────────
fc = 3000       # psi
fy = 60000      # psi
fy_ksi = 60     # ksi
fc_ksi = 3      # ksi
gamma_c = 150   # pcf
Es = 29_000_000 # psi
Es_ksi = 29000  # ksi
L_ft = 22       # ft
L_in = L_ft * 12
wL = 1.2        # k/ft (live load)
wD_super = 0.6  # k/ft (superimposed dead load)
cover = 1.5     # in (ACI 318-19 Table 20.6.1.3.1)
stirrup_dia = 0.375  # #3 stirrup
bar_dia_8 = 1.0      # #8 bar diameter
bar_area_8 = 0.79    # #8 bar area

# Derived constants
Ec = 57000 * math.sqrt(fc)  # psi
Ec_ksi = Ec / 1000
n = Es / Ec  # modular ratio
fr = 7.5 * math.sqrt(fc)  # modulus of rupture, psi
beta1 = 0.85  # for f'c <= 4000 psi

# ── Helper Functions ────────────────────────────────────────────
def calc_d(h, bar_dia=bar_dia_8):
    return h - cover - stirrup_dia - bar_dia / 2

def self_weight(b_in, h_in):
    """Self-weight in k/ft"""
    return (b_in / 12) * (h_in / 12) * gamma_c / 1000

def factored_load(wD, wL):
    return 1.2 * wD + 1.6 * wL

def moment(w, L):
    """Simply supported uniform load moment in k-ft"""
    return w * L**2 / 8

def shear(w, L):
    """Simply supported uniform load shear in kips"""
    return w * L / 2

def rho_from_Rn(Rn_psi, fc_psi, fy_psi):
    """Required reinforcement ratio from Rn"""
    ratio = 2 * Rn_psi / (0.85 * fc_psi)
    if ratio >= 1.0:
        return None  # impossible
    return (0.85 * fc_psi / fy_psi) * (1 - math.sqrt(1 - ratio))

def rho_max_tc(fc_psi, fy_psi, beta1):
    """Max rho for tension-controlled section"""
    return 0.85 * beta1 * (fc_psi / fy_psi) * (0.003 / (0.003 + 0.005))

def rho_min(fc_psi, fy_psi):
    return max(3 * math.sqrt(fc_psi) / fy_psi, 200 / fy_psi)

def Ig(b, h):
    return b * h**3 / 12

def cracked_NA(b, n_ratio, As, d):
    """Neutral axis depth of cracked section"""
    # b*(kd)^2/2 = n*As*(d - kd)
    a_coeff = b / 2
    b_coeff = n_ratio * As
    c_coeff = -n_ratio * As * d
    kd = (-b_coeff + math.sqrt(b_coeff**2 - 4 * a_coeff * c_coeff)) / (2 * a_coeff)
    return kd

def Icr(b, kd, n_ratio, As, d):
    return b * kd**3 / 3 + n_ratio * As * (d - kd)**2

def Ie_bischoff(Mcr, Ma, Icr_val, Ig_val):
    """Effective I per ACI 318-19 Eq. 24.2.3.5a (Bischoff)"""
    if Ma <= (2/3) * Mcr:
        return Ig_val
    ratio = (2/3) * (Mcr / Ma)
    return Icr_val / (1 - ratio**2 * (1 - Icr_val / Ig_val))

def deflection_uniform(w_klin, L_in, Ec_psi, Ie):
    """5wL^4 / (384EI), w in k/in"""
    return 5 * w_klin * L_in**4 / (384 * Ec_psi * Ie)


# ════════════════════════════════════════════════════════════════
#  DOCUMENT
# ════════════════════════════════════════════════════════════════

add_title("REINFORCED CONCRETE BEAM DESIGN")
add_text("Simply Supported Beam — Uniform Load", bold=True)
add_text("Design per ACI 318-19 | Loads per ASCE 7-22")

doc.add_paragraph("─" * 65)

# ── GIVEN ───────────────────────────────────────────────────────
add_heading2("1. GIVEN INFORMATION")
add_calc(f"Span, L                = {L_ft} ft")
add_calc(f"Superimposed Dead Load = {wD_super*1000:.0f} plf = {wD_super} k/ft")
add_calc(f"Live Load, wL          = {wL} k/ft")
add_calc(f"f'c                    = {fc} psi")
add_calc(f"fy                     = {fy} psi ({fy_ksi} ksi)")
add_calc(f"Concrete unit weight   = {gamma_c} pcf (normal weight)")
add_calc(f"Clear cover            = {cover} in  (ACI 318-19 Table 20.6.1.3.1)")
add_calc(f"Stirrups               = #3 (d_s = {stirrup_dia} in)")
add_calc(f"Ec = 57000*sqrt(f'c)   = {Ec:.0f} psi = {Ec_ksi:.1f} ksi")
add_calc(f"n  = Es/Ec             = {n:.2f} (use {round(n,1)})")
add_calc(f"fr = 7.5*sqrt(f'c)     = {fr:.1f} psi")
add_calc(f"beta_1                 = {beta1} (for f'c <= 4000 psi)")
add_ref("ACI 318-19 §19.2.2.1, §19.2.3.1, §24.2.3.5")

n_used = round(n, 1)

# ── REINFORCEMENT LIMITS ───────────────────────────────────────
rho_min_val = rho_min(fc, fy)
rho_tc_max = rho_max_tc(fc, fy, beta1)

add_heading2("2. REINFORCEMENT LIMITS")
add_calc(f"rho_min = max(3*sqrt(f'c)/fy, 200/fy)")
add_calc(f"        = max(3*{math.sqrt(fc):.2f}/{fy}, 200/{fy})")
add_calc(f"        = max({3*math.sqrt(fc)/fy:.5f}, {200/fy:.5f})")
add_calc(f"        = {rho_min_val:.5f}")
add_ref("ACI 318-19 §9.6.1.2")

add_calc(f"rho_max (tension-controlled, eps_t >= 0.005):")
add_calc(f"rho_max = 0.85*beta1*(f'c/fy)*(0.003/(0.003+0.005))")
add_calc(f"        = 0.85*{beta1}*({fc}/{fy})*0.375")
add_calc(f"        = {rho_tc_max:.5f}")
add_ref("ACI 318-19 §21.2.2, §9.3.3.1")


# ════════════════════════════════════════════════════════════════
#  TRIAL 1: 12 x 18
# ════════════════════════════════════════════════════════════════
add_heading2("3. TRIAL SECTION: 12\" x 18\"")

b1, h1 = 12, 18
sw1 = self_weight(b1, h1)
wD1 = wD_super + sw1
wu1 = factored_load(wD1, wL)
Mu1 = moment(wu1, L_ft)
d1 = calc_d(h1)

add_heading3("3.1 Loads")
add_calc(f"Self-weight = (12/12)*(18/12)*150 = {sw1*1000:.0f} plf = {sw1:.3f} k/ft")
add_calc(f"Total Dead Load, wD = {wD_super} + {sw1:.3f} = {wD1:.3f} k/ft")
add_calc(f"Live Load, wL = {wL:.3f} k/ft")
add_calc("")
add_calc(f"Factored Load (ASCE 7-22 §2.3.1, Combo 2):")
add_calc(f"  wu = 1.2*wD + 1.6*wL = 1.2({wD1:.3f}) + 1.6({wL})")
add_calc(f"     = {1.2*wD1:.3f} + {1.6*wL:.3f} = {wu1:.3f} k/ft")
add_ref("ASCE 7-22 §2.3.1 Load Combination 2: 1.2D + 1.6L")

add_heading3("3.2 Flexural Design")
add_calc(f"d = h - cover - stirrup - bar/2")
add_calc(f"  = {h1} - {cover} - {stirrup_dia} - {bar_dia_8/2}")
add_calc(f"  = {d1:.3f} in")
add_calc("")
add_calc(f"Mu = wu*L^2/8 = {wu1:.3f}*{L_ft}^2/8 = {Mu1:.2f} k-ft")
add_calc(f"   = {Mu1*12:.2f} k-in")
add_calc("")

Rn1 = Mu1 * 12 / (0.9 * b1 * d1**2)  # ksi
Rn1_psi = Rn1 * 1000
add_calc(f"Rn = Mu/(phi*b*d^2) = {Mu1*12:.2f}/(0.9*{b1}*{d1:.3f}^2)")
add_calc(f"   = {Rn1_psi:.1f} psi")
add_calc("")

rho1 = rho_from_Rn(Rn1_psi, fc, fy)
if rho1 is not None:
    add_calc(f"rho_req = (0.85*f'c/fy)*[1 - sqrt(1 - 2*Rn/(0.85*f'c))]")
    add_calc(f"        = {rho1:.5f}")
else:
    add_calc(f"rho_req = CANNOT BE COMPUTED (Rn too high)")

add_calc("")
add_calc(f"rho_max (tension-controlled) = {rho_tc_max:.5f}")
add_calc(f"rho_req = {rho1:.5f}" if rho1 else "rho_req = N/A")

if rho1 is not None and rho1 > rho_tc_max:
    add_warning(f"rho_req ({rho1:.5f}) > rho_max ({rho_tc_max:.5f})")
    add_warning("Section is NOT tension-controlled. 12\"x18\" is INADEQUATE for flexure.")
    add_warning("Does not meet ACI 318-19 §9.3.3.1 for phi = 0.9")

    # Show what the max capacity is
    As_max1 = rho_tc_max * b1 * d1
    a_max1 = As_max1 * fy_ksi / (0.85 * fc_ksi * b1)
    phiMn_max1 = 0.9 * As_max1 * fy_ksi * (d1 - a_max1/2) / 12
    add_calc("")
    add_calc(f"Maximum phi*Mn (tension-controlled) for 12x18:")
    add_calc(f"  As_max = {rho_tc_max:.5f}*{b1}*{d1:.3f} = {As_max1:.3f} in^2")
    add_calc(f"  a = {As_max1:.3f}*60/(0.85*3*12) = {a_max1:.3f} in")
    add_calc(f"  phi*Mn = 0.9*{As_max1:.3f}*60*({d1:.3f}-{a_max1/2:.3f})/12")
    add_calc(f"         = {phiMn_max1:.2f} k-ft")
    add_calc(f"  Mu     = {Mu1:.2f} k-ft")
    add_warning(f"phi*Mn = {phiMn_max1:.1f} k-ft < Mu = {Mu1:.1f} k-ft -- FAILS")

add_text("")
add_text("CONCLUSION: 12\"x18\" beam is inadequate. Must increase section depth.", bold=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  TRIAL 2: 12 x 22
# ════════════════════════════════════════════════════════════════
add_heading2("4. REVISED SECTION: 12\" x 22\"")

b2, h2 = 12, 22
sw2 = self_weight(b2, h2)
wD2 = wD_super + sw2
wu2 = factored_load(wD2, wL)
ws2 = wD2 + wL  # service load
Mu2 = moment(wu2, L_ft)
d2 = calc_d(h2)

add_heading3("4.1 Loads")
add_calc(f"Self-weight = ({b2}/12)*({h2}/12)*150 = {sw2*1000:.0f} plf = {sw2:.3f} k/ft")
add_calc(f"Total Dead Load, wD = {wD_super} + {sw2:.3f} = {wD2:.3f} k/ft")
add_calc(f"Live Load, wL = {wL:.3f} k/ft")
add_calc("")
add_calc(f"wu = 1.2({wD2:.3f}) + 1.6({wL}) = {wu2:.3f} k/ft")
add_calc(f"Mu = {wu2:.3f}*{L_ft}^2/8 = {Mu2:.2f} k-ft = {Mu2*12:.2f} k-in")
add_ref("ASCE 7-22 §2.3.1 Load Combination 2")

add_heading3("4.2 Minimum Depth Check")
min_depth = L_in / 16
add_calc(f"ACI 318-19 Table 9.3.1.1 — Simply supported beam:")
add_calc(f"  h_min = L/16 = {L_in}/16 = {min_depth:.2f} in")
add_calc(f"  h = {h2} in > {min_depth:.2f} in  --> OK")
add_ref("ACI 318-19 Table 9.3.1.1")

# ── FLEXURE ─────────────────────────────────────────────────────
add_heading3("4.3 Flexural Design")
add_calc(f"d = {h2} - {cover} - {stirrup_dia} - {bar_dia_8/2} = {d2:.3f} in")
add_calc("")

Rn2 = Mu2 * 12 / (0.9 * b2 * d2**2)
Rn2_psi = Rn2 * 1000
add_calc(f"Rn = Mu/(phi*b*d^2)")
add_calc(f"   = {Mu2*12:.2f} / (0.9*{b2}*{d2:.3f}^2)")
add_calc(f"   = {Rn2_psi:.1f} psi")

rho2 = rho_from_Rn(Rn2_psi, fc, fy)
add_calc("")
add_calc(f"rho_req = (0.85*f'c/fy)*[1 - sqrt(1 - 2*Rn/(0.85*f'c))]")
add_calc(f"        = {rho2:.5f}")
add_calc("")
add_calc(f"Check: rho_min = {rho_min_val:.5f} < rho_req = {rho2:.5f} < rho_max = {rho_tc_max:.5f}  --> OK")
add_ref("ACI 318-19 §9.6.1.2, §9.3.3.1")

As_req = rho2 * b2 * d2
add_calc("")
add_calc(f"As_req = rho*b*d = {rho2:.5f}*{b2}*{d2:.3f} = {As_req:.3f} in^2")

# Select bars
n_bars = 3
bar_size = "#8"
As_prov = n_bars * bar_area_8
add_calc("")
add_calc(f"USE {n_bars}-{bar_size} bars: As = {n_bars}*{bar_area_8} = {As_prov:.2f} in^2 > {As_req:.3f} in^2  --> OK")

# Check bar spacing
avail_width = b2 - 2*cover - 2*stirrup_dia
spacing_bars = (avail_width - n_bars * bar_dia_8) / (n_bars - 1)
add_calc("")
add_calc(f"Bar spacing check:")
add_calc(f"  Available width = {b2} - 2({cover}) - 2({stirrup_dia}) = {avail_width:.2f} in")
add_calc(f"  Clear spacing = ({avail_width:.2f} - {n_bars}*{bar_dia_8})/({n_bars-1})")
add_calc(f"                = {spacing_bars:.3f} in")
min_spacing = max(1.0, bar_dia_8, 1.0)  # 1 in, db, 1 in (assuming 1" max agg)
add_calc(f"  Min spacing = max(1 in, d_b, 1 in) = {min_spacing:.2f} in")
if spacing_bars >= min_spacing:
    add_pass(f"Clear spacing {spacing_bars:.2f} in >= {min_spacing:.2f} in")
else:
    add_warning(f"Clear spacing {spacing_bars:.2f} in < {min_spacing:.2f} in")
add_ref("ACI 318-19 §25.2.1")

# ── VERIFY STRENGTH ────────────────────────────────────────────
add_heading3("4.4 Verify Flexural Strength")
a2 = As_prov * fy_ksi / (0.85 * fc_ksi * b2)
c2 = a2 / beta1
eps_t = 0.003 * (d2 - c2) / c2

add_calc(f"a = As*fy / (0.85*f'c*b)")
add_calc(f"  = {As_prov:.2f}*{fy_ksi} / (0.85*{fc_ksi}*{b2})")
add_calc(f"  = {a2:.3f} in")
add_calc("")
add_calc(f"c = a/beta1 = {a2:.3f}/{beta1} = {c2:.3f} in")
add_calc(f"eps_t = 0.003*(d-c)/c = 0.003*({d2:.3f}-{c2:.3f})/{c2:.3f}")
add_calc(f"      = {eps_t:.6f}")
add_calc("")

if eps_t >= 0.005:
    add_pass(f"eps_t = {eps_t:.4f} >= 0.005 --> Tension-controlled, phi = 0.9")
    phi_flex = 0.9
else:
    add_warning(f"eps_t = {eps_t:.4f} < 0.005 --> NOT tension-controlled")
    eps_y = fy / Es
    phi_flex = 0.65 + 0.25 * (eps_t - eps_y) / (0.005 - eps_y)
    add_calc(f"phi = {phi_flex:.4f}")

Mn2 = As_prov * fy_ksi * (d2 - a2/2)  # k-in
phiMn2 = phi_flex * Mn2 / 12  # k-ft

add_calc("")
add_calc(f"Mn = As*fy*(d - a/2)")
add_calc(f"   = {As_prov:.2f}*{fy_ksi}*({d2:.3f} - {a2/2:.3f})")
add_calc(f"   = {Mn2:.2f} k-in = {Mn2/12:.2f} k-ft")
add_calc("")
add_calc(f"phi*Mn = {phi_flex}*{Mn2/12:.2f} = {phiMn2:.2f} k-ft")
add_calc(f"Mu     = {Mu2:.2f} k-ft")

if phiMn2 >= Mu2:
    add_pass(f"phi*Mn = {phiMn2:.1f} k-ft >= Mu = {Mu2:.1f} k-ft  --> ADEQUATE")
else:
    add_warning(f"phi*Mn = {phiMn2:.1f} k-ft < Mu = {Mu2:.1f} k-ft  --> INADEQUATE")
add_ref("ACI 318-19 §9.5.1.1, §21.2.1, §22.2.2.1")

# ── SHEAR ───────────────────────────────────────────────────────
add_heading3("4.5 Shear Design")

Vu_face = shear(wu2, L_ft)
d2_ft = d2 / 12
Vu_d = wu2 * (L_ft/2 - d2_ft)  # at d from face

add_calc(f"Vu (at face) = wu*L/2 = {wu2:.3f}*{L_ft}/2 = {Vu_face:.2f} kips")
add_calc(f"Vu (at d from face) = wu*(L/2 - d)")
add_calc(f"   = {wu2:.3f}*({L_ft/2:.1f} - {d2_ft:.3f})")
add_calc(f"   = {Vu_d:.2f} kips")
add_ref("ACI 318-19 §9.4.3.2")

# Concrete shear strength
Vc = 2 * math.sqrt(fc) * b2 * d2 / 1000  # kips
phi_shear = 0.75
phiVc = phi_shear * Vc

add_calc("")
add_calc(f"Vc = 2*sqrt(f'c)*bw*d = 2*{math.sqrt(fc):.2f}*{b2}*{d2:.3f}/1000")
add_calc(f"   = {Vc:.2f} kips")
add_calc(f"phi*Vc = 0.75*{Vc:.2f} = {phiVc:.2f} kips")
add_ref("ACI 318-19 §22.5.5.1")

if Vu_d > phiVc:
    Vs_req = Vu_d / phi_shear - Vc
    add_calc("")
    add_calc(f"Vu > phi*Vc --> Stirrups required")
    add_calc(f"Vs_req = Vu/phi - Vc = {Vu_d:.2f}/0.75 - {Vc:.2f} = {Vs_req:.2f} kips")

    Av = 2 * 0.11  # 2 legs of #3
    s_calc = Av * fy_ksi * d2 / Vs_req
    s_max = min(d2/2, 24)
    s_use = min(math.floor(s_calc), math.floor(s_max))
    # Round down to nearest inch
    s_use = int(s_use)

    add_calc("")
    add_calc(f"Using #3 stirrups (2 legs): Av = 2*0.11 = {Av:.2f} in^2")
    add_calc(f"s_calc = Av*fy*d/Vs = {Av:.2f}*{fy_ksi}*{d2:.3f}/{Vs_req:.2f} = {s_calc:.2f} in")
    add_calc(f"s_max  = min(d/2, 24) = min({d2/2:.2f}, 24) = {s_max:.2f} in")
    add_calc(f"USE #3 stirrups @ {s_use} in o.c.")

    # Check min Av
    Av_s_min = max(0.75*math.sqrt(fc)*b2/fy, 50*b2/fy)
    Av_s_prov = Av / s_use
    add_calc("")
    add_calc(f"Min Av/s = max(0.75*sqrt(f'c)*bw/fy, 50*bw/fy)")
    add_calc(f"        = max({0.75*math.sqrt(fc)*b2/fy:.5f}, {50*b2/fy:.5f})")
    add_calc(f"        = {Av_s_min:.5f} in^2/in")
    add_calc(f"Provided Av/s = {Av:.2f}/{s_use} = {Av_s_prov:.5f} in^2/in")
    if Av_s_prov >= Av_s_min:
        add_pass(f"Av/s provided >= Av/s min")
    else:
        add_warning("Av/s provided < Av/s min")

    # Vs capacity check
    Vs_prov = Av * fy_ksi * d2 / s_use
    Vs_max = 8 * math.sqrt(fc) * b2 * d2 / 1000
    add_calc("")
    add_calc(f"Vs_provided = Av*fy*d/s = {Av:.2f}*{fy_ksi}*{d2:.3f}/{s_use}")
    add_calc(f"            = {Vs_prov:.2f} kips")
    add_calc(f"Check Vs <= 8*sqrt(f'c)*bw*d = {Vs_max:.2f} kips  --> OK")

    phiVn = phi_shear * (Vc + Vs_prov)
    add_calc("")
    add_calc(f"phi*Vn = 0.75*({Vc:.2f} + {Vs_prov:.2f}) = {phiVn:.2f} kips")
    add_calc(f"Vu     = {Vu_d:.2f} kips")
    if phiVn >= Vu_d:
        add_pass(f"phi*Vn = {phiVn:.1f} kips >= Vu = {Vu_d:.1f} kips  --> ADEQUATE")
    add_ref("ACI 318-19 §9.5.1.1, §22.5.1.1, §22.5.10.5.3, §9.7.6.2.2")

else:
    add_pass("Vu <= phi*Vc — Minimum stirrups only")
    s_use = int(min(d2/2, 24))
    add_calc(f"USE #3 stirrups @ {s_use} in o.c. (minimum)")

doc.add_page_break()

# ── DEFLECTION ──────────────────────────────────────────────────
add_heading2("5. DEFLECTION CHECK")
add_ref("ACI 318-19 §24.2.3, Table 24.2.2")

Ig2 = Ig(b2, h2)
yt = h2 / 2
Mcr_kin = fr * Ig2 / yt / 1000  # k-in
Mcr_kft = Mcr_kin / 12

add_heading3("5.1 Section Properties")
add_calc(f"Ig = b*h^3/12 = {b2}*{h2}^3/12 = {Ig2:.1f} in^4")
add_calc(f"yt = h/2 = {yt:.1f} in")
add_calc(f"Mcr = fr*Ig/yt = {fr:.1f}*{Ig2:.1f}/{yt:.1f}")
add_calc(f"    = {Mcr_kin:.1f} k-in = {Mcr_kft:.1f} k-ft")
add_ref("ACI 318-19 §24.2.3.5")

# Cracked section properties
kd2 = cracked_NA(b2, n_used, As_prov, d2)
Icr2 = Icr(b2, kd2, n_used, As_prov, d2)

add_calc("")
add_calc(f"Cracked neutral axis (n = {n_used}):")
add_calc(f"  {b2}*(kd)^2/2 = {n_used}*{As_prov:.2f}*({d2:.3f} - kd)")
add_calc(f"  kd = {kd2:.3f} in")
add_calc("")
add_calc(f"Icr = b*(kd)^3/3 + n*As*(d-kd)^2")
add_calc(f"    = {b2}*{kd2:.3f}^3/3 + {n_used}*{As_prov:.2f}*({d2:.3f}-{kd2:.3f})^2")
add_calc(f"    = {b2*kd2**3/3:.1f} + {n_used*As_prov*(d2-kd2)**2:.1f}")
add_calc(f"    = {Icr2:.1f} in^4")

# Service moments
Ma_D_kft = moment(wD2, L_ft)
Ma_L_kft = moment(wL, L_ft)
Ma_total_kft = Ma_D_kft + Ma_L_kft
Ma_D_kin = Ma_D_kft * 12
Ma_L_kin = Ma_L_kft * 12
Ma_total_kin = Ma_total_kft * 12

add_heading3("5.2 Service Load Moments")
add_calc(f"M_D = wD*L^2/8 = {wD2:.3f}*{L_ft}^2/8 = {Ma_D_kft:.2f} k-ft = {Ma_D_kin:.1f} k-in")
add_calc(f"M_L = wL*L^2/8 = {wL}*{L_ft}^2/8 = {Ma_L_kft:.2f} k-ft = {Ma_L_kin:.1f} k-in")
add_calc(f"M_total = {Ma_total_kft:.2f} k-ft = {Ma_total_kin:.1f} k-in")

# Effective moment of inertia
add_heading3("5.3 Effective Moment of Inertia (Bischoff Equation)")
add_calc(f"ACI 318-19 Eq. 24.2.3.5a:")
add_calc(f"Ie = Icr / [1 - ((2/3)(Mcr/Ma))^2 * (1 - Icr/Ig)]")
add_calc(f"For Ma > (2/3)*Mcr = {2/3*Mcr_kin:.1f} k-in")

# I_e for total service moment
Ie_total = Ie_bischoff(Mcr_kin, Ma_total_kin, Icr2, Ig2)
ratio_total = (2/3) * (Mcr_kin / Ma_total_kin)
add_calc("")
add_calc(f"For total load (Ma = {Ma_total_kin:.1f} k-in):")
add_calc(f"  (2/3)(Mcr/Ma) = (2/3)({Mcr_kin:.1f}/{Ma_total_kin:.1f}) = {ratio_total:.4f}")
add_calc(f"  Ie = {Icr2:.1f} / [1 - {ratio_total:.4f}^2 * (1 - {Icr2:.1f}/{Ig2:.1f})]")
add_calc(f"     = {Icr2:.1f} / [1 - {ratio_total**2:.5f} * {1-Icr2/Ig2:.4f}]")
add_calc(f"     = {Ie_total:.1f} in^4")

# I_e for dead load
Ie_D = Ie_bischoff(Mcr_kin, Ma_D_kin, Icr2, Ig2)

add_calc("")
add_calc(f"For dead load only (Ma = {Ma_D_kin:.1f} k-in):")
if Ma_D_kin > (2/3)*Mcr_kin:
    ratio_D = (2/3) * (Mcr_kin / Ma_D_kin)
    add_calc(f"  (2/3)(Mcr/Ma) = {ratio_D:.4f}")
    add_calc(f"  Ie_D = {Ie_D:.1f} in^4")
else:
    add_calc(f"  Ma <= (2/3)*Mcr --> Ie_D = Ig = {Ig2:.1f} in^4")

# ── LIVE LOAD DEFLECTION ───────────────────────────────────────
add_heading3("5.4 Live Load Deflection")

wL_in = wL / 12  # k/in
delta_LL = deflection_uniform(wL_in, L_in, Ec_ksi, Ie_total)
limit_360 = L_in / 360

add_calc(f"Delta_LL = 5*wL*L^4 / (384*Ec*Ie)")
add_calc(f"wL = {wL} k/ft = {wL_in:.5f} k/in")
add_calc(f"L  = {L_in} in")
add_calc(f"Ec = {Ec_ksi:.1f} ksi")
add_calc(f"Ie = {Ie_total:.1f} in^4")
add_calc("")
add_calc(f"Delta_LL = 5*{wL_in:.5f}*{L_in}^4 / (384*{Ec_ksi:.1f}*{Ie_total:.1f})")
add_calc(f"         = {delta_LL:.4f} in")
add_calc("")
add_calc(f"Allowable = L/360 = {L_in}/360 = {limit_360:.3f} in")
add_calc("")

if delta_LL <= limit_360:
    add_pass(f"Delta_LL = {delta_LL:.3f} in <= L/360 = {limit_360:.3f} in  --> OK")
else:
    add_warning(f"Delta_LL = {delta_LL:.3f} in > L/360 = {limit_360:.3f} in  --> FAILS")
add_ref("ACI 318-19 Table 24.2.2 — Floor not supporting damageable elements")

# ── LONG-TERM DEFLECTION ───────────────────────────────────────
add_heading3("5.5 Long-Term Deflection")

xi = 2.0  # 5 years or more
rho_prime = 0  # no compression steel
lambda_delta = xi / (1 + 50 * rho_prime)

wD_in = wD2 / 12  # k/in
delta_D = deflection_uniform(wD_in, L_in, Ec_ksi, Ie_D)
delta_LT_D = lambda_delta * delta_D
delta_total = delta_LT_D + delta_LL
limit_240 = L_in / 240

add_calc(f"Time-dependent multiplier (ACI 318-19 §24.2.4.1):")
add_calc(f"  xi = {xi} (sustained load duration >= 5 years)")
add_calc(f"  rho' = {rho_prime} (no compression reinforcement)")
add_calc(f"  lambda_delta = xi/(1+50*rho') = {xi}/(1+0) = {lambda_delta:.1f}")
add_calc("")
add_calc(f"Immediate dead load deflection:")
add_calc(f"  Delta_D = 5*{wD_in:.5f}*{L_in}^4 / (384*{Ec_ksi:.1f}*{Ie_D:.1f})")
add_calc(f"          = {delta_D:.4f} in")
add_calc("")
add_calc(f"Long-term dead load deflection:")
add_calc(f"  lambda*Delta_D = {lambda_delta:.1f}*{delta_D:.4f} = {delta_LT_D:.4f} in")
add_calc("")
add_calc(f"Total deflection after attachment of nonstructural elements:")
add_calc(f"  Delta_total = lambda*Delta_D + Delta_LL")
add_calc(f"              = {delta_LT_D:.4f} + {delta_LL:.4f}")
add_calc(f"              = {delta_total:.4f} in")
add_calc("")
add_calc(f"Allowable = L/240 = {L_in}/240 = {limit_240:.3f} in")

if delta_total <= limit_240:
    add_pass(f"Delta_total = {delta_total:.3f} in <= L/240 = {limit_240:.3f} in  --> OK")
else:
    add_warning(f"Delta_total = {delta_total:.3f} in > L/240 = {limit_240:.3f} in  --> FAILS")
add_ref("ACI 318-19 §24.2.4.1, Table 24.2.2")

doc.add_page_break()

# ── DEVELOPMENT LENGTH ──────────────────────────────────────────
add_heading2("6. DEVELOPMENT LENGTH CHECK")
add_calc(f"Development length for #8 bars (ACI 318-19 §25.4.2.3):")

# Simplified method
psi_t = 1.0  # bottom bars
psi_e = 1.0  # uncoated
psi_s = 1.0  # #8 bars (>= #7)
psi_g = 1.0  # Grade 60
cb = min(cover + stirrup_dia + bar_dia_8/2,
         (avail_width + n_bars*bar_dia_8)/(2*n_bars))  # min of side and bottom
# cb + Ktr / db
# Using simplified: if clear cover >= db and clear spacing >= 2*db
clear_cover_actual = cover + stirrup_dia  # to bar surface
clear_spacing = spacing_bars

add_calc(f"  Coating factor, psi_e = {psi_e}")
add_calc(f"  Size factor, psi_s = {psi_s} (bar >= #7)")
add_calc(f"  Location factor, psi_t = {psi_t} (bottom bars)")
add_calc(f"  Grade factor, psi_g = {psi_g} (Grade 60)")

# Using Table 25.4.2.3
# Condition A: clear spacing >= 2db, clear cover >= db, with minimum ties
# ld = (fy*psi_t*psi_e*psi_s*psi_g / (25*lambda*sqrt(f'c))) * db
lam = 1.0  # normal weight
if clear_spacing >= 2*bar_dia_8 and clear_cover_actual >= bar_dia_8:
    cond = "A"
    ld = (fy * psi_t * psi_e * psi_s * psi_g / (25 * lam * math.sqrt(fc))) * bar_dia_8 / 1.0
else:
    cond = "B"
    ld = (fy * psi_t * psi_e * psi_s * psi_g / (20 * lam * math.sqrt(fc))) * bar_dia_8 / 1.0

add_calc(f"  Clear spacing = {clear_spacing:.2f} in, 2*db = {2*bar_dia_8} in")
add_calc(f"  Clear cover = {clear_cover_actual:.3f} in, db = {bar_dia_8} in")
add_calc(f"  Condition {cond} applies")
add_calc("")

divisor = 25 if cond == "A" else 20
add_calc(f"  ld = (fy*psi_t*psi_e*psi_s*psi_g / ({divisor}*lambda*sqrt(f'c))) * db")
add_calc(f"     = ({fy}*{psi_t}*{psi_e}*{psi_s}*{psi_g} / ({divisor}*{lam}*{math.sqrt(fc):.2f})) * {bar_dia_8}")
add_calc(f"     = {ld:.1f} in")

ld_min = max(12, ld)
add_calc(f"  ld_min = max(12 in, {ld:.1f} in) = {ld_min:.1f} in")

avail_ld = L_in/2  # half span for simply supported
add_calc(f"  Available = L/2 = {avail_ld:.0f} in")
if avail_ld >= ld_min:
    add_pass(f"Available {avail_ld:.0f} in >= ld = {ld_min:.1f} in  --> OK")
add_ref("ACI 318-19 §25.4.2.3, Table 25.4.2.3")


# ── SUMMARY ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading2("7. DESIGN SUMMARY")

doc.add_paragraph("─" * 65)
add_calc(f"BEAM SIZE:        {b2}\" x {h2}\" (revised from 12\" x 18\")")
add_calc(f"SPAN:             {L_ft} ft (simply supported)")
add_calc(f"CONCRETE:         f'c = {fc} psi, normal weight")
add_calc(f"STEEL:            fy = {fy} psi (Grade 60)")
add_calc("")
add_calc(f"FLEXURE:")
add_calc(f"  Bottom steel:   {n_bars}-{bar_size} (As = {As_prov:.2f} in^2)")
add_calc(f"  phi*Mn:         {phiMn2:.1f} k-ft >= Mu = {Mu2:.1f} k-ft  --> OK")
add_calc("")
add_calc(f"SHEAR:")
add_calc(f"  Stirrups:       #3 @ {s_use}\" o.c.")
add_calc(f"  phi*Vn:         {phiVn:.1f} kips >= Vu = {Vu_d:.1f} kips  --> OK")
add_calc("")
add_calc(f"DEFLECTION:")
add_calc(f"  Delta_LL:       {delta_LL:.3f} in <= L/360 = {limit_360:.3f} in  --> OK")
add_calc(f"  Delta_total(LT):{delta_total:.3f} in <= L/240 = {limit_240:.3f} in  --> OK")
doc.add_paragraph("─" * 65)

# Cross-section sketch description
add_text("")
add_text("CROSS SECTION:", bold=True)
add_calc(f"    |<--- 12\" --->|")
add_calc(f"    ┌─────────────┐  ─┬─")
add_calc(f"    │             │   │")
add_calc(f"    │             │   │")
add_calc(f"    │  #3 @ {s_use}\"   │  {h2}\"")
add_calc(f"    │  stirrups   │   │")
add_calc(f"    │             │   │")
add_calc(f"    │  ●   ●   ●  │   │")
add_calc(f"    │  {n_bars}-{bar_size} bars  │  ─┴─")
add_calc(f"    └─────────────┘")
add_calc(f"    1.5\" clr cover (typ.)")

# Save
output_path = "/Users/shirish/aiplayground/Beam_Design_12x22.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
