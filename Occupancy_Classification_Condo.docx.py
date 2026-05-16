"""
Occupancy Classification for a 16-Story Residential Condominium in Florida
Generates a .docx calculation/reference sheet
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_title(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(14)

def add_heading2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(12)

def add_heading3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(11)

def add_calc(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"NOTE: {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 180)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_ref(text):
    p = doc.add_paragraph()
    run = p.add_run(f"Reference: {text}")
    run.italic = True
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    return p

# ════════════════════════════════════════════════════════════════
#  DOCUMENT
# ════════════════════════════════════════════════════════════════

add_title("OCCUPANCY & CONSTRUCTION TYPE CLASSIFICATION")
add_text("16-Story Residential Condominium — Florida", bold=True)
add_text("Per Florida Building Code (FBC) 8th Edition & ASCE 7-22")
doc.add_paragraph("─" * 65)

# ── SECTION 1: OCCUPANCY GROUP ─────────────────────────────────
add_heading2("1. OCCUPANCY GROUP CLASSIFICATION")

add_text("Topic: Occupancy Group for Residential Condominium", bold=True)
add_text("")

add_text("Applicable Code: FBC 8th Ed., Chapter 3, §302.1", bold=True)
add_text("")

add_calc("Code Requirement:")
add_calc("  FBC 8th Ed. §302.1 — Occupancy Classification")
add_calc("  Table 302.1 classifies buildings by their use or character of occupancy.")
add_calc("")
add_calc("  Group R — Residential:")
add_calc("  Residential Group R includes, among others, the use of a")
add_calc("   building or structure, or a portion thereof, for sleeping")
add_calc("   purposes when not classified as an Institutional Group I")
add_calc("   or when not regulated by the Florida Fire Prevention Code.")
add_calc("")
add_calc("  Group R is further divided into:")
add_calc("    R-1: Transient (hotels, motels, boarding houses)")
add_calc("    R-2: Permanent (apartments, condominiums, dormitories)")
add_calc("    R-3: One- and two-family dwellings")
add_calc("    R-4: Assisted living (5-16 persons)")
add_ref("FBC 8th Ed. §302.1, Table 302.1")

add_calc("")
add_calc("Application:")
add_calc("─" * 50)
add_calc("  Given:  16-story residential condominium")
add_calc("  Find:   Occupancy Group")
add_calc("")
add_calc("  A condominium is a permanent residential occupancy where")
add_calc("  occupants are primarily permanent (non-transient) residents.")
add_calc("")
add_calc("  Per FBC 8th Ed. §310.1.1 — Residential Group R-2:")
add_calc("  Residential occupancies containing sleeping units or more")
add_calc("   than two dwelling units where the occupants are primarily")
add_calc("   permanent in nature, including:")
add_calc("     - Apartment houses")
add_calc("     - Boarding houses (non-transient) with more than 16 occupants")
add_calc("     - Convents")
add_calc("     - Dormitories")
add_calc("     - Fraternities and sororities")
add_calc("     - Hotels (non-transient)")
add_calc("     - Live/work units")
add_calc("     - Monasteries")
add_calc("     - Vacation timeshare properties")
add_calc("")
add_calc("  ┌─────────────────────────────────────────────────┐")
add_calc("  │  RESULT: OCCUPANCY GROUP R-2                    │")
add_calc("  │  (Residential — Permanent, > 2 dwelling units)  │")
add_calc("  └─────────────────────────────────────────────────┘")
add_ref("FBC 8th Ed. §310.1.1, Table 302.1")

# ── SECTION 2: RISK CATEGORY ──────────────────────────────────
add_heading2("2. RISK CATEGORY")

add_text("Applicable Code: ASCE 7-22, Table 1.5-1", bold=True)
add_calc("")
add_calc("ASCE 7-22 Table 1.5-1 — Risk Category of Buildings:")
add_calc("")
add_calc("  Risk Category I:   Low risk to human life")
add_calc("  Risk Category II:  All buildings not in I, III, or IV")
add_calc("  Risk Category III: Buildings with large occupant loads (>300)")
add_calc("                     or potential for substantial economic/mass")
add_calc("                     disruption")
add_calc("  Risk Category IV:  Essential facilities")
add_calc("")
add_calc("  For a 16-story condominium:")
add_calc("  - Residential Group R-2")
add_calc("  - Not an essential facility")
add_calc("  - Must check occupant load threshold:")
add_calc("")
add_calc("    Typical condo floor: assume 8-12 units/floor")
add_calc("    16 floors x ~10 units x ~2.5 persons = ~400 persons")
add_calc("")
add_calc("    Per ASCE 7-22 Table 1.5-1, Footnote:")
add_calc("    Buildings with an occupant load greater than 300 at any")
add_calc("     time are Risk Category III, HOWEVER this applies to")
add_calc("     assembly and educational uses, not residential.")
add_calc("")
add_calc("    Residential buildings (Group R) are classified as:")

add_calc("")
add_calc("  ┌─────────────────────────────────────────────────┐")
add_calc("  │  RESULT: RISK CATEGORY II                       │")
add_calc("  │  (Standard residential occupancy)               │")
add_calc("  └─────────────────────────────────────────────────┘")
add_ref("ASCE 7-22 §1.5, Table 1.5-1")

add_note("If the condominium includes ground-floor assembly spaces (lobby, "
         "clubhouse, restaurant) with occupant loads > 300, those portions "
         "may trigger Risk Category III. Verify with the occupant load "
         "calculation per FBC §1004.5.")

# ── SECTION 3: CONSTRUCTION TYPE ───────────────────────────────
add_heading2("3. CONSTRUCTION TYPE CONSIDERATIONS")

add_text("Applicable Code: FBC 8th Ed., Chapter 6, Tables 601 & 602", bold=True)
add_calc("")
add_calc("For a 16-story building, the construction type is constrained")
add_calc("by height and area limits per FBC 8th Ed. Table 504.4.")
add_calc("")
add_calc("  Group R-2, 16 stories:")
add_calc("  ─────────────────────────────────────────────────")
add_calc("  Type I-A:  Unlimited height, Unlimited area      --> OK")
add_calc("  Type I-B:  12 stories / 180 ft (with sprinklers) --> MARGINAL")
add_calc("  Type II-A: 5 stories max                         --> NO")
add_calc("  Type II-B: 4 stories max                         --> NO")
add_calc("  Type III:  4-5 stories max                       --> NO")
add_calc("  Type IV:   4-5 stories max                       --> NO")
add_calc("  Type V:    3-4 stories max                       --> NO")
add_calc("")
add_calc("  Note: Story limits above include sprinkler increases")
add_calc("        per FBC 8th Ed. §504.2")
add_calc("")
add_calc("  ┌─────────────────────────────────────────────────┐")
add_calc("  │  REQUIRED: TYPE I-A CONSTRUCTION                │")
add_calc("  │  (Noncombustible, highest fire ratings)         │")
add_calc("  └─────────────────────────────────────────────────┘")
add_calc("")
add_calc("  Type I-A Fire Resistance Ratings (FBC Table 601):")
add_calc("  ─────────────────────────────────────────────────")
add_calc("    Structural Frame:             3 hours")
add_calc("    Bearing Walls (exterior):     3 hours")
add_calc("    Bearing Walls (interior):     3 hours")
add_calc("    Floor Construction:           2 hours")
add_calc("    Roof Construction:            1.5 hours")
add_ref("FBC 8th Ed. §602.2, Table 504.4, Table 601")

# ── SECTION 4: LIVE LOADS ──────────────────────────────────────
add_heading2("4. APPLICABLE LIVE LOADS")

add_text("Applicable Code: ASCE 7-22, Table 4.3-1", bold=True)
add_calc("")
add_calc("  Minimum Live Loads for Residential (R-2):")
add_calc("  ─────────────────────────────────────────────────")
add_calc("  Private rooms & corridors serving them:  40 psf")
add_calc("  Public rooms & corridors serving them:   100 psf")
add_calc("  Corridors above first floor:             80 psf*")
add_calc("  Lobbies/first floor corridors:           100 psf")
add_calc("  Stairs and exits:                        100 psf")
add_calc("  Balconies (exterior):                    60 psf or 1.5x")
add_calc("  Parking garages:                         40 psf")
add_calc("  Mechanical/electrical rooms:             125 psf**")
add_calc("  Roof (ordinary):                         20 psf")
add_calc("")
add_calc("  * Per ASCE 7-22 Table 4.3-1, corridors above first")
add_calc("    floor follow the occupancy they serve")
add_calc("  ** Or actual equipment weight, whichever is greater")
add_ref("ASCE 7-22 Table 4.3-1")

# ── SUMMARY ─────────────────────────────────────────────────────
doc.add_page_break()
add_heading2("5. SUMMARY")

doc.add_paragraph("─" * 65)
add_calc("  BUILDING:           16-Story Residential Condominium")
add_calc("                      Located in Florida")
add_calc("")
add_calc("  OCCUPANCY GROUP:    R-2 (Residential, permanent)")
add_calc("                      FBC 8th Ed. §310.1.1, Table 302.1")
add_calc("")
add_calc("  RISK CATEGORY:      II (Standard occupancy)")
add_calc("                      ASCE 7-22 Table 1.5-1")
add_calc("")
add_calc("  CONSTRUCTION TYPE:  Type I-A (Noncombustible)")
add_calc("                      FBC 8th Ed. Table 504.4, §602.2")
add_calc("")
add_calc("  FIRE RATINGS:       Structural frame: 3 hr")
add_calc("                      Bearing walls:    3 hr")
add_calc("                      Floor:            2 hr")
add_calc("                      Roof:             1.5 hr")
add_calc("                      FBC 8th Ed. Table 601")
add_calc("")
add_calc("  RESIDENTIAL LL:     40 psf (private units)")
add_calc("                      100 psf (public corridors/lobby)")
add_calc("                      ASCE 7-22 Table 4.3-1")
doc.add_paragraph("─" * 65)

# Save
output_path = "/Users/shirish/aiplayground/Occupancy_Classification_16Story_Condo.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
