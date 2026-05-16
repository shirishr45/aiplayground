#!/usr/bin/env python3
"""Convert ClaudeCodeStudy.md to DOCX format"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_docx(md_file, docx_file):
    doc = Document()

    # Read markdown
    with open(md_file, 'r') as f:
        content = f.read()

    # Remove code block markers for cleaner output
    lines = content.split('\n')

    current_code_block = []
    in_code_block = False

    for line in lines:
        # Skip table of contents links (they don't work in DOCX anyway)
        if line.strip().startswith('[') and '](#' in line:
            continue

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # Add accumulated code
                if current_code_block:
                    p = doc.add_paragraph()
                    p.add_run(' '.join(current_code_block)).font.name = 'Consolas'
                current_code_block = []
                in_code_block = not in_code_block
            continue

        if in_code_block:
            current_code_block.append(line)
            continue

        # Handle headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # Handle tables
        elif line.strip().startswith('|---') or line.strip().startswith('|---'):
            continue  # Skip separator lines
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            # Table row - add as simple text
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and any(cells):
                doc.add_paragraph(' | '.join(cells))

        # Handle bullet points
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')

        # Handle numbered lists
        elif line.strip().startswith('### Exercise') or line.strip().startswith('### Week'):
            doc.add_heading(line.strip(), level=3)

        # Handle regular text (skip empty lines and frontmatter)
        elif line.strip() and not line.startswith('>') and not line.startswith('---'):
            p = doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Saved: {docx_file}")

if __name__ == "__main__":
    parse_markdown_to_docx(
        '/Users/shirish/aiplayground/ClaudeCodeStudy.md',
        '/Users/shirish/aiplayground/ClaudeCodeStudy.docx'
    )
